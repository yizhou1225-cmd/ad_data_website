import streamlit as st
import pandas as pd
import numpy as np
import io
import json
import xlsxwriter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import docx.opc.constants
import time

COMMON_METRICS = {
    "spend": ["已花费金额 (USD)", "花费金额（USD）"],

    "conversion": ["成效","潜在客户人数"],
    "cpe": ["单次成效费用"],

    "cpm": ["千次展示费用"],
    "cpc": ["单次链接点击费用","单次点击费用（全部）"],

    "impressions": ["展示次数", "曝光"],

    # ✅ 正确归类
    "ctr_all": ["点击率（全部）", "链接点击率（%）"],

    # ✅ 新增（非常重要）
    "reach": ["覆盖人数"],
    "frequency": ["频次"],
}
SHEET_MAPPINGS = {
    "整体数据": {
        **COMMON_METRICS,
    },

    "分时段数据": {
        **COMMON_METRICS,
        "date": ["单日"]
    },

    "国家": {
        **COMMON_METRICS, 
        "dimension_item": ["国家/地区"]
    },

    "性别": {
        **COMMON_METRICS,
        "dimension_item": ["性别"]
    },

    "年龄": {
        **COMMON_METRICS,
        "dimension_item": ["年龄"]
    },

    "平台&版位": {
        **COMMON_METRICS,
        "dimension_item": ["版位"]
    },

    "受众组": {
        **COMMON_METRICS,
        "dimension_item": ["广告组名称", "广告系列名称"],
        "custom_audience_settings": ["设置的自定义受众", "Custom Audiences"],
        "converting_keywords": ["产生成效的关键词", "Interests", "Keywords"],
        "converting_countries": ["产生成效的国家", "国家", "地区", "Country", "Region", "Location"],
        "converting_genders": ["产生成效的性别", "性别", "Gender"],
        "converting_ages": ["产生成效的年龄", "年龄", "Age", "Age Group"]
    },

    # ✅ clicks 只在素材
    "素材": {
        **COMMON_METRICS,
        "clicks": ["点击"],
        "content_item": ["素材"]
    }
}

FIELD_ALIASES = {
    "spend": ["spend", "amount spent", "花费"],
    "conversion": ["成效", "results", "潜在客户人数"],

    "impressions": ["impressions", "展示", "曝光"],

    "ctr_all": ["点击率（全部）", "ctr (all)"],

    "cpc": ["cpc", "单次链接点击费用","单次点击费用（全部）"],
    "cpm": ["cpm", "千次展示费用"],
    "cpe": ["cpe", "单次成效费用"],
    "converting_countries": ["converting_countries", "country", "region", "国家", "地区", "location"],
    "converting_genders": ["converting_genders", "gender", "性别"],
    "converting_ages": ["converting_ages", "age", "年龄", "age_group"],
    "converting_keywords": ["converting_keywords", "keywords", "interests", "兴趣", "关键词"],
    # ✅ clicks只留一个语义
    "clicks": ["点击"],

    "dimension_item": ["国家/地区", "受众组", "性别", "年龄", "版位"],
    "content_item": ["素材"]
}

GROUP_CONFIG = {
    "Master_Overview": ["整体数据", "分时段数据"],
    "Master_Breakdown": ["国家","受众组", "受众类型", "性别", "年龄", "平台&版位"],
    "Master_Creative": ["素材"]
}

# ✅ 👇 放这里（关键）
REPORT_MAPPING = {
    "spend": "花费 ($)",
    "conversion": "成效",
    "cpe": "单次成效费用 ($)",
    "impressions": "展现量",
    "ctr_all": "点击率 (All)",
    "cpm": "CPM",
    "cpc": "CPC",
    "reach": "覆盖人数",
    "frequency": "频次",
    "date_range": "日期",
    "dimension_item": "维度",
    "content_item": "素材名称",
    "converting_countries": "产生成效的国家", "converting_genders": "产生成效的性别", "converting_ages": "产生成效的年龄"
}

# ==========================================
# PART 2: 核心工具函数 (已修复百分比识别问题)
# ==========================================

def parse_float(value):
    """辅助函数：清理数据并将字符串/数字安全转换为浮点数"""
    if value is None:
        return 0.0
    try:
        if isinstance(value, (int, float)):
            return float(value)
        return clean_numeric_strict(value)
    except (ValueError, TypeError):
        return 0.0

def safe_div(numerator, denominator, multiplier=1.0):
    n = parse_float(numerator)
    d = parse_float(denominator)
    if d > 0:
        return (n / d) * multiplier
    else:
        return 0.0

def clean_numeric(val):
    # ✅ 防止 Series / list
    if isinstance(val, (pd.Series, list, np.ndarray)):
        return 0.0

    if pd.isna(val):
        return 0.0

    if isinstance(val, (int, float)):
        return float(val)

    val_str = str(val).strip().replace('$', '').replace('¥', '').replace(',', '')

    if '%' in val_str:
        val_str = val_str.replace('%', '')
        try:
            return float(val_str) / 100.0
        except:
            return 0.0

    # ✅ ❗必须有这一段（你缺的关键）
    try:
        return float(val_str)
    except:
        return 0.0

def clean_numeric_strict(val):
    # ✅ 防炸核心
    if isinstance(val, (pd.Series, list, np.ndarray)):
        return 0.0

    if pd.isna(val):
        return 0.0

    if isinstance(val, (int, float)):
        return float(val)

    val_str = str(val).strip().replace('$', '').replace('¥', '').replace(',', '')

    if '%' in val_str:
        val_str = val_str.replace('%', '')
        try:
            return float(val_str) / 100.0
        except:
            return 0.0

    try:
        return float(val_str)
    except:
        return 0.0

def find_column_fuzzy(df, keywords):
    cols = [str(c) for c in df.columns]

    for kw in keywords:
        kw_norm = str(kw).lower().replace(" ", "").replace("_", "")

        for col in cols:
            col_norm = col.lower().replace(" ", "").replace("_", "")

            # ✅ 优先完全匹配
            if kw_norm == col_norm:
                return col

        for col in cols:
            col_norm = col.lower().replace(" ", "").replace("_", "")

            # ✅ 再用包含匹配（但更安全）
            if kw_norm in col_norm:
                return col

    return None

def calc_metrics_dict(df_chunk):
    res = {}
    if df_chunk.empty:
        return res

    # ========================
    # 1️⃣ 只汇总“量”
    # ========================
    for col in ['spend', 'clicks', 'impressions', 'conversion']:
        found = find_column_fuzzy(df_chunk, FIELD_ALIASES.get(col, [col]))
        if found:
            res[col] = df_chunk[found].apply(clean_numeric_strict).sum()
        else:
            res[col] = 0.0

    # ========================
    # 2️⃣ 指标：直接用原字段（不计算）
    # ========================
    for col in ['ctr', 'cpc', 'cpm', 'cpe']:
        found = find_column_fuzzy(df_chunk, FIELD_ALIASES.get(col, [col]))
        if found:
            # 👉 用平均（或者你也可以用 .iloc[0]）
            res[col] = df_chunk[found].apply(clean_numeric_strict).mean()
        else:
            res[col] = 0.0

    # ========================
    # 日期处理（稳定版）
    # ========================
    if 'date' in df_chunk.columns:
        try:
            dates = pd.to_datetime(df_chunk['date'], errors='coerce').dropna()
            if not dates.empty:
                res['date_range'] = f"{dates.min():%Y-%m-%d} ~ {dates.max():%Y-%m-%d}"
            else:
                res['date_range'] = "-"
        except:
            res['date_range'] = "-"

    elif 'date_range' in df_chunk.columns:
        res['date_range'] = str(df_chunk['date_range'].iloc[0])

    else:
        res['date_range'] = "-"

    return res

def format_cell(key, val, is_mom=False):
    if isinstance(val, str): return val
    if is_mom:
        if key == 'date_range': return val
        return f"{val:+.2%}"
    k = str(key).lower()
    if any(x in k for x in ['rate', 'ctr', 'cvr', '点击率', '转化率', '着陆率', '意向率', '成功率']): 
        return f"{val:.2%}" 
    if any(x in k for x in ['spend', 'cpm', 'cpc', 'value', '花费', '金额', '客单价', 'gmv', '价值']): return f"{val:,.2f}"
    if any(x in k for x in ['purchases', 'cart', 'click', '次数', '单量', '点击', '展现', '访问量', '发起数']): return f"{val:,.0f}"
    return f"{val}"

def extract_benchmark_values(df_bench):
    targets = {'cpm': (['cpm'], False), 'ctr': (['ctr'], True), 'cpc': (['cpc'], False)}
    extracted = {}
    for metric, (aliases, higher_better) in targets.items():
        found_col = None
        for alias in aliases:
            found_col = find_column_fuzzy(df_bench, [alias])
            if found_col: break
        if found_col:
            try:
                s = df_bench[found_col].apply(clean_numeric_strict)
                v = s[s>0].mean()
                if metric in ['ctr'] and v > 1.0:
                    v = v / 100.0
                if not pd.isna(v): extracted[metric] = [v, higher_better]
            except: pass
    return extracted

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    try:
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color)
            rPr.append(c)
        if underline:
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink
    except: return None

def apply_report_labels(df, custom_mapping=None):
    if df.empty: return df
    mapping = REPORT_MAPPING.copy()
    if custom_mapping: mapping.update(custom_mapping)
    return df.rename(columns=mapping)

def add_df_to_word(doc, df, title, level=1):
    if df.empty: return
    doc.add_heading(title, level=level)
    t = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    t.style = 'Table Grid'
    is_creative = "素材" in title
    is_landing = "落地页" in title
    link_col_idx = -1
    for j, col in enumerate(df.columns):
        cell = t.cell(0, j)
        cell.text = str(col)
        if any(x in str(col).lower() for x in ["url", "link", "素材", "内容", "content"]): link_col_idx = j
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(8)
    for i in range(df.shape[0]):
        label_prefix = "素材" if is_creative else ("落地页" if is_landing else "")
        label_char = chr(65 + (i % 26))
        if i >= 26: label_char += str(i // 26)
        label_text = f"{label_prefix}{label_char}"
        for j in range(df.shape[1]):
            val = df.iat[i, j]
            cell = t.cell(i+1, j)
            if (is_creative or is_landing) and j == link_col_idx:
                try:
                    p = cell.paragraphs[0]
                    url = str(val).strip()
                    if len(url) > 5: add_hyperlink(p, url, label_text)
                    else: cell.text = label_text
                except: cell.text = label_text
            else:
                cell.text = str(val)
                if "结论" in str(df.columns[j]):
                    if "✅" in str(val): cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                    if "⚠️" in str(val): cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(8)
    doc.add_paragraph("\n")

# ==========================================
# PART 3: 主逻辑类
# ==========================================

class AdReportProcessor:
    def __init__(self, raw_files, bench_file=None):
        self.raw_files = raw_files
        self.bench_file = bench_file
        self.processed_dfs = {}
        self.merged_dfs = {}
        self.final_json = {}
        self.doc = Document()

    def detect_sheet_type(self, df, file_name=""):
        cols = [str(c).strip().lower() for c in df.columns]
        file_name = str(file_name).lower()

        def has_col(keyword_list):
            for c in cols:
                for kw in keyword_list:
                    if kw.lower() in c:
                        return True
            return False

        # 先按文件名判断（最稳）
        if "单日" in file_name:
            return "分时段数据"
        # ✅ 新增
        if "总表" in file_name:
            return "整体数据"
        if "目标" in file_name:
            return "受众组"
        if "国家" in file_name:
            return "国家"
        if "性别" in file_name:
            return "性别"
        if "年龄" in file_name:
            return "年龄"
        if "版位" in file_name:
            return "平台&版位"
        if "素材" in file_name:
            return "素材"
        if "广告系列" in file_name:
            return "受众组"

        # 再按字段判断（兜底）
        if has_col(["单日"]):
            return "分时段数据"
        if has_col(["国家/地区", "国家"]):
            return "国家"
        if has_col(["性别"]):
            return "性别"
        if has_col(["年龄"]):
            return "年龄"
        if has_col(["版位"]):
            return "平台&版位"
        if has_col(["素材"]):
            return "素材"
        if has_col(["广告组名称", "广告系列名称"]):
            return "受众组"

        return None

    def process_etl(self):
        for file in self.raw_files:
            try:
                xls = pd.ExcelFile(file)
                if not xls.sheet_names:
                    continue

                # 默认读取第一个非空sheet
                df = None
                for sheet in xls.sheet_names:
                    temp_df = pd.read_excel(file, sheet_name=sheet)
                    if not temp_df.empty:
                        df = temp_df
                        break

                if df is None or df.empty:
                    continue


                matched_sheet = self.detect_sheet_type(df, file.name)

                if not matched_sheet:
                    st.warning(f"⚠️ 无法识别文件类型，已跳过：{file.name}")
                    continue

                mapping = SHEET_MAPPINGS.get(matched_sheet, {})
                final_cols = {}

                for std_col, raw_col_options in mapping.items():
                    print(f"\n🧠 正在匹配字段: {std_col}")
                    print(f"候选字段: {raw_col_options}")
                    print(f"当前Excel列: {df.columns.tolist()}")
                    matched_col = find_column_fuzzy(df, raw_col_options)
                    
                    if matched_col:
                        print(f"✅ 匹配成功: {std_col} -> {matched_col}")
                        final_cols[std_col] = matched_col
                    else:
                        print(f"❌ 匹配失败: {std_col}")

                if not final_cols:
                    st.warning(f"⚠️ 文件已识别但没有匹配到有效字段：{file.name}")
                    continue

                unique_cols = list(dict.fromkeys(final_cols.values()))

                df_clean = df[unique_cols].rename(
                    columns={v: k for k, v in final_cols.items()}
                )
                print(f"\n📦 清洗后字段（{matched_sheet}）:")
                print(df_clean.columns.tolist())

                # 再去重 rename 后列
                df_clean = df_clean.loc[:, ~df_clean.columns.duplicated()]

                text_cols = [
                    'date','date_range', 'anomaly_metric_name',
                    'converting_keywords', 'converting_countries',
                    'converting_genders', 'converting_ages',
                    'custom_audience_settings', 'dimension_item', 'content_item'
                ]

                for col in df_clean.columns:
                    if col not in text_cols:
                        df_clean[col] = df_clean[col].apply(clean_numeric)
                
                # ✅ 自动补CPE（只在缺失时）
                if "cpe" not in df_clean.columns:
                    if "spend" in df_clean.columns and "conversion" in df_clean.columns:
                        print(f"⚙️ 自动计算 cpe（{matched_sheet}）")
                        df_clean["cpe"] = (
                            df_clean["spend"] / df_clean["conversion"].replace(0, None)
                        )

                # ✅ ✅ ✅ 就放在这里（非常关键）
                if "dimension_item" in df_clean.columns:
                    df_clean["dimension_item"] = df_clean["dimension_item"].astype(str)
                
                # 保留来源
                df_clean["Source_Sheet"] = matched_sheet

                if matched_sheet == "整体数据":
                    start_col = find_column_fuzzy(df, ["报告开始日期", "开始日期"])
                    end_col = find_column_fuzzy(df, ["报告结束日期", "结束日期"])

                    if start_col and end_col:
                        df_clean["date_range"] = (
                            df[start_col].astype(str).str.strip() +
                            " ~ " +
                            df[end_col].astype(str).str.strip()
                        )

                # 某些表只保留前10
                if matched_sheet in ["素材", "落地页", "受众组"]:
                    if "spend" in df_clean.columns:
                        df_clean = df_clean.sort_values("spend", ascending=False).head(10)

                if matched_sheet not in self.processed_dfs:
                    self.processed_dfs[matched_sheet] = []

                self.processed_dfs[matched_sheet].append(df_clean)

            except Exception as e:
                st.warning(f"⚠️ 文件处理失败：{file.name}，原因：{e}")

        # 合并同类文件
        for sheet_name, df_list in self.processed_dfs.items():
            if isinstance(df_list, list) and df_list:
                self.processed_dfs[sheet_name] = pd.concat(df_list, ignore_index=True)

        # 生成 merged_dfs（保留你原来的逻辑）
        for master_name, source_sheets in GROUP_CONFIG.items():
            dfs_to_merge = [self.processed_dfs[src] for src in source_sheets if src in self.processed_dfs]
            if dfs_to_merge:
                merged_df = pd.concat(dfs_to_merge, ignore_index=True)
                cols = list(merged_df.columns)
                priority_cols = ['Source_Sheet', 'date_range', 'dimension_item', 'content_item',
                                'spend']
                new_order = [c for c in priority_cols if c in cols] + [c for c in cols if c not in priority_cols]
                self.merged_dfs[master_name] = merged_df[new_order]

    def generate_report(self):
        benchmark_targets = {'cpm': [20.0, False], 'ctr': [0.015, True], 'cpc': [1.5, False]}
        if self.bench_file:
            try:
                df_b = pd.read_excel(self.bench_file)
                benchmark_targets = extract_benchmark_values(df_b)
            except:
                pass

        self.doc.add_heading('广告投放深度分析报告', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.final_json = {
            "report_title": "广告投放深度分析报告",
            "generated_at": pd.Timestamp.now().strftime("%Y-%m-%d")
        }

        # 1. 大盘总览
        df_ov = pd.DataFrame()
        if "Master_Overview" in self.merged_dfs:
            df_src = self.merged_dfs["Master_Overview"]
            mask = df_src['Source_Sheet'].astype(str).apply(lambda x: any(k in x for k in ["分时", "Time"]))
            df_ov = df_src[mask].copy() if not df_src[mask].empty else df_src.copy()

        if not df_ov.empty:
            date_col = find_column_fuzzy(df_ov, ['date', 'time', '时间'])
            if date_col:
                try:
                    df_ov['temp_date'] = pd.to_datetime(df_ov[date_col], errors='coerce')
                    df_clean = df_ov.dropna(subset=['temp_date']).sort_values('temp_date')

                    # ✅ 关键1：日期排序
                    dates = np.sort(df_clean['temp_date'].dropna().unique())

                    raw_overall = calc_metrics_dict(df_clean)

                    if len(dates) >= 2:
                        # ✅ 关键2：正确切分时间
                        mid_date = dates[len(dates)//2]

                        df_prev = df_clean[df_clean['temp_date'] < mid_date]
                        df_curr = df_clean[df_clean['temp_date'] >= mid_date]

                        # 👉 debug（建议保留一轮）
                        print("前半周期行数:", len(df_prev))
                        print("后半周期行数:", len(df_curr))

                        raw_prev = calc_metrics_dict(df_prev)
                        raw_curr = calc_metrics_dict(df_curr)

                        # ✅ 关键3：环比计算（安全版）
                        raw_mom = {}
                        for k in raw_curr:
                            if k == 'date_range':
                                raw_mom[k] = "-"
                                continue

                            v_prev = raw_prev.get(k, 0)
                            v_curr = raw_curr.get(k, 0)

                            # 👉 防止除0
                            if isinstance(v_prev, (int, float)) and v_prev != 0:
                                raw_mom[k] = (v_curr - v_prev) / v_prev
                            else:
                                raw_mom[k] = 0.0

                    else:
                        # 数据不足兜底
                        raw_prev = {k: "-" for k in raw_overall}
                        raw_curr = raw_overall
                        raw_mom = {k: "-" for k in raw_overall}

        # =========================
        # 输出结构
        # =========================

                    col_order = ["date_range", "spend", "cpe", "cpm", "cpc", "ctr"]
                    final_data = []
                    for label, r in zip(["整体数据", "前半周期", "后半周期", "环比"], [raw_overall, raw_prev, raw_curr, raw_mom]):
                        row = {"Label": label}
                        is_m = (label == "环比")
                        for c in col_order:
                            row[c] = format_cell(c, r.get(c, 0), is_mom=is_m)
                        row['date_range'] = label
                        final_data.append(row)

                    df_f = pd.DataFrame(final_data, columns=col_order)
                    df_f_display = apply_report_labels(df_f)
                    add_df_to_word(self.doc, df_f_display, "1. 数据大盘总览", level=1)
                    self.final_json['1_data_overview'] = df_f_display.to_dict(orient='records')

                    raw_current = calc_metrics_dict(df_clean)
                    bench_data = []
                    for metric_key in ['cpe', 'cpm', 'ctr', 'cpc']:
                        curr_val = raw_current.get(metric_key, 0)
                        bench_val, higher_is_better = benchmark_targets.get(metric_key, [0, True])
                        conclusion = "-"
                        if curr_val != 0:
                            diff = curr_val - bench_val
                            if higher_is_better:
                                conclusion = "✅ 优于大盘" if diff > 0 else ("⚠️ 低于大盘" if diff < 0 else "持平")
                            else:
                                conclusion = "✅ 优于大盘" if diff < 0 else ("⚠️ 高于大盘" if diff > 0 else "持平")

                        bench_data.append({
                            "指标": REPORT_MAPPING.get(metric_key, metric_key.upper()),
                            "当前账户": format_cell(metric_key, curr_val),
                            "行业基准": format_cell(metric_key, bench_val),
                            "对比结论": conclusion
                        })

                    df_b = pd.DataFrame(bench_data)
                    add_df_to_word(self.doc, df_b, "2. 行业 Benchmark 对比", level=1)
                    self.final_json['2_industry_benchmark'] = df_b.to_dict(orient='records')

                except Exception as e:
                    st.warning(f"大盘计算警告: {e}")

        # 3. 受众分析
        self.doc.add_heading("3. 受众分析", level=1)
        self.final_json['3_audience_analysis'] = {}

        def resolve_std_col(df, std_col):
            print(f"\n🔍 resolve_std_col: 正在找 {std_col}")
            print("当前列:", df.columns.tolist())

            if std_col in df.columns:
                print(f"✅ 直接命中: {std_col}")
                return std_col

            aliases = FIELD_ALIASES.get(std_col, [std_col])
            col = find_column_fuzzy(df, aliases)

            if col:
                print(f"✅ 模糊匹配成功: {std_col} -> {col}")
            else:
                print(f"❌ 匹配失败: {std_col}")

            return col
            aliases = FIELD_ALIASES.get(std_col, [std_col])
            return find_column_fuzzy(df, aliases)

        def build_audience_table(df_curr, title, dim_label, top10=False, include_converting=False, level=2):
            if df_curr.empty:
                return

            df_curr = df_curr.copy()
            df_curr.columns = df_curr.columns.astype(str)

            for metric in ['ctr', 'cpc', 'cpm']:
                col = resolve_std_col(df_curr, metric)
                df_curr[metric] = df_curr[col] if col else np.nan

            req_cols = ["dimension_item", "spend", "conversion", "cpe", "ctr", "cpc", "cpm"]

            if include_converting:
                req_cols += [
                    "custom_audience_settings",
                    "converting_countries",
                    "converting_keywords",
                    "converting_genders",
                    "converting_ages"
                ]

            rename_map = {}
            valid_cols = []

            for req in req_cols:
                found = resolve_std_col(df_curr, req)

                if found:
                    valid_cols.append(found)
                    if found != req:
                        rename_map[found] = req
                else:
                    default_val = "-" if ("converting" in req or req == "custom_audience_settings") else np.nan
                    df_curr[req] = default_val
                    valid_cols.append(req)

            df_final = df_curr[valid_cols].rename(columns=rename_map)

            for t_col in [
                "custom_audience_settings",
                "converting_countries",
                "converting_keywords",
                "converting_genders",
                "converting_ages"
            ]:
                if t_col in df_final.columns:
                    df_final[t_col] = df_final[t_col].fillna("-").astype(str).replace("nan", "-")

            if "dimension_item" in df_final.columns:
                df_final = df_final[
                    ~df_final['dimension_item'].astype(str).str.lower().str.contains('unknow', na=False)
                ]

            if top10 and 'spend' in df_final.columns:
                df_final = df_final.sort_values('spend', ascending=False).head(10)

            df_clean = df_final.round(2)
            df_display = apply_report_labels(df_clean, custom_mapping={'dimension_item': dim_label})

            add_df_to_word(self.doc, df_display, title, level=level)
            self.final_json['3_audience_analysis'][title] = df_display.to_dict(orient='records')

        if "Master_Breakdown" in self.merged_dfs:
            df_bd = self.merged_dfs["Master_Breakdown"]

            def safe_filter(df, keyword):
                return df[df["Source_Sheet"].astype(str).str.contains(keyword, na=False)].copy()

            df_country = safe_filter(df_bd, "国家")
            build_audience_table(df_country, "3.1 国家分析", "国家", top10=True, include_converting=False, level=2)

            df_gender = safe_filter(df_bd, "性别")
            build_audience_table(df_gender, "3.2 性别分析", "性别", top10=False, include_converting=False, level=2)

            df_age = safe_filter(df_bd, "年龄")
            build_audience_table(df_age, "3.3 年龄分析", "年龄段", top10=False, include_converting=False, level=2)

            df_adset = safe_filter(df_bd, "受众组")
            build_audience_table(df_adset, "3.4 受众组分析表", "受众组名称", top10=True, include_converting=True, level=2)

            df_audtype = safe_filter(df_bd, "受众类型")
            build_audience_table(df_audtype, "3.5 受众类型分析", "受众类型", top10=False, include_converting=False, level=2)

        # 4. 素材
        if "Master_Creative" in self.merged_dfs:
            df_cr = self.merged_dfs["Master_Creative"]
            mask = df_cr['Source_Sheet'].astype(str).str.contains("素材", na=False)
            df_curr = df_cr[mask].copy()

            if not df_curr.empty:
                df_curr = df_curr.copy()
                df_curr.columns = df_curr.columns.astype(str)

                for metric in ['ctr', 'cpc', 'cpm']:
                    col = resolve_std_col(df_curr, metric)
                    df_curr[metric] = df_curr[col] if col else np.nan

                req_cols = ["content_item", "spend", "conversion", "cpe", "ctr", "cpc", "cpm"]

                rename_map = {}
                valid_cols = []

                for req in req_cols:
                    found = resolve_std_col(df_curr, req)

                    if found:
                        valid_cols.append(found)
                        if found != req:
                            rename_map[found] = req
                    else:
                        df_curr[req] = np.nan
                        valid_cols.append(req)

                df_final = df_curr[valid_cols].rename(columns=rename_map)

                if 'spend' in df_final.columns:
                    df_final = df_final.sort_values('spend', ascending=False).head(10)

                df_clean = df_final.round(2)
                df_display = apply_report_labels(df_clean, custom_mapping={'content_item': "素材名称"})

                add_df_to_word(self.doc, df_display, "4. 素材分析", level=1)
                self.final_json["4_creative_analysis"] = df_display.to_dict(orient='records')

        # 5. 版位
        if "Master_Breakdown" in self.merged_dfs:
            self.doc.add_heading("5. 版位分析", level=1)

            df_bd = self.merged_dfs["Master_Breakdown"]
            mask = df_bd['Source_Sheet'].astype(str).str.contains("版位|Placement", na=False, regex=True)
            df_curr = df_bd[mask].copy()

            if not df_curr.empty:
                df_curr = df_curr.copy()
                df_curr.columns = df_curr.columns.astype(str)

                for metric in ['ctr', 'cpc', 'cpm']:
                    col = resolve_std_col(df_curr, metric)
                    df_curr[metric] = df_curr[col] if col else np.nan

                req_cols = ['dimension_item', "spend", "conversion", "cpe", "ctr", "cpc", "cpm"]

                rename_map = {}
                valid_cols = []

                for c in req_cols:
                    f = resolve_std_col(df_curr, c)

                    if f:
                        valid_cols.append(f)
                        if f != c:
                            rename_map[f] = c
                    else:
                        df_curr[c] = np.nan
                        valid_cols.append(c)

                df_clean = df_curr[valid_cols].rename(columns=rename_map).round(2)

                df_top5 = df_clean.sort_values('spend', ascending=False).head(5)
                add_df_to_word(self.doc, apply_report_labels(df_top5, {'dimension_item': '版位'}), "5.1 版位花费 TOP 5", level=2)

                if 'ctr' in df_clean.columns and 'cpm' in df_clean.columns:
                    mean_ctr = df_clean['ctr'].mean()
                    mean_cpm = df_clean['cpm'].mean()

                    df_pot = df_clean[
                        (df_clean['ctr'] > mean_ctr) & (df_clean['cpm'] < mean_cpm)
                    ].sort_values('ctr', ascending=False).head(5)

                    if df_pot.empty:
                        df_pot = df_clean.sort_values('ctr', ascending=False).head(5)
                else:
                    df_pot = df_clean.head(5)

                add_df_to_word(self.doc, apply_report_labels(df_pot, {'dimension_item': '版位'}), "5.2 版位高潜力", level=2)

                self.final_json['5_placement_analysis'] = {
                    "top_spend": apply_report_labels(df_top5, {'dimension_item': '版位'}).to_dict('records'),
                    "high_potential": apply_report_labels(df_pot, {'dimension_item': '版位'}).to_dict('records')
                }

# ==========================================
# PART 4: Streamlit UI
# ==========================================

def main():
    st.set_page_config(page_title="Auto-ad-data", layout="wide")

    st.markdown("""
        <style>
        .stApp {
            background: radial-gradient(circle at 50% 20%, rgba(255, 240, 200, 0.6) 0%, rgba(240, 200, 255, 0.4) 30%, rgba(255, 255, 255, 1) 70%);
            background-attachment: fixed;
            background-size: cover;
        }
        .main-title {
            text-align: center;
            font-size: 3.5rem;
            font-weight: 800;
            margin-bottom: 0.5rem;
            font-family: -apple-system, BlinkMacSystemFont, sans-serif;
            background: linear-gradient(135deg, #662D8C 0%, #ED1E79 100%);
            background: -webkit-linear-gradient(135deg, #662D8C 0%, #ED1E79 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }
        .sub-title {
            text-align: center;
            font-size: 1.1rem;
            color: #666;
            margin-bottom: 3rem;
        }
        [data-testid="stVerticalBlockBorderWrapper"] {
            background: rgba(255, 255, 255, 0.45) !important;
            backdrop-filter: blur(14px);
            -webkit-backdrop-filter: blur(14px);
            border: 1px solid rgba(255, 255, 255, 0.8) !important;
            box-shadow: 0 10px 30px rgba(0, 0, 0, 0.04);
            border-radius: 24px !important;
            padding: 1rem;
            transition: all 0.3s ease;
        }
        [data-testid="stVerticalBlockBorderWrapper"]:hover {
            transform: translateY(-3px);
            box-shadow: 0 15px 35px rgba(102, 45, 140, 0.1);
            border-color: rgba(237, 30, 121, 0.2) !important;
            background: rgba(255, 255, 255, 0.65) !important;
        }
        .card-header {
            text-align: center;
            font-weight: 600;
            color: #4A4A4A;
            margin-bottom: 0.5rem;
            letter-spacing: 0.5px;
        }
        .icon-container {
            text-align: center;
            font-size: 2.5rem;
            margin-bottom: 5px;
            filter: drop-shadow(0 2px 4px rgba(0,0,0,0.1));
        }
        [data-testid='stFileUploader'] section {
            background-color: rgba(255, 255, 255, 0.25);
            backdrop-filter: blur(15px);
            -webkit-backdrop-filter: blur(15px);
            border: 1.5px dashed rgba(102, 45, 140, 0.3);
            box-shadow: inset 0 0 20px rgba(255, 255, 255, 0.5);
            border-radius: 16px;
            padding: 1.5rem;
            transition: all 0.3s ease;
        }
        [data-testid='stFileUploader'] section:hover {
            background-color: rgba(255, 255, 255, 0.5);
            border-color: #ED1E79;
            box-shadow: 0 8px 20px rgba(102, 45, 140, 0.15);
        }
        [data-testid='stFileUploader'] button {
            border-radius: 20px;
            border-color: rgba(102, 45, 140, 0.2);
            color: #662D8C;
            background-color: rgba(255, 255, 255, 0.8);
        }
        .glass-info-box {
            padding: 1rem 1.5rem;
            margin-bottom: 1.5rem;
            display: flex;
            align-items: center;
            border-radius: 16px;
            background: rgba(255, 255, 255, 0.3);
            backdrop-filter: blur(15px);
            -webkit-backdrop-filter: blur(15px);
            border: 1px solid rgba(255, 255, 255, 0.7);
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.05);
        }
        div.stButton > button {
            display: block;
            margin: 0 auto;
            width: 100%;
            background-image: linear-gradient(90deg, #B721FF 0%, #21D4FD 100%);
            color: white !important;
            border-radius: 30px;
            padding: 0.7rem 1.5rem;
            font-size: 1.1rem;
            font-weight: 600;
            border: none;
            box-shadow: 0 6px 20px rgba(183, 33, 255, 0.4); 
            transition: all 0.3s ease;
        }
        div.stButton > button:hover {
            opacity: 0.95;
            transform: translateY(-3px) scale(1.02);
            box-shadow: 0 10px 30px rgba(33, 212, 253, 0.5);
        }
        div.stDownloadButton > button {
            display: block;
            margin: 0 auto;
            width: 100%;
            background: linear-gradient(145deg, rgba(255, 255, 255, 1) 0%, rgba(245, 235, 255, 1) 100%);
            color: #662D8C !important; 
            border-radius: 30px;
            padding: 0.7rem 1.5rem;
            font-size: 1rem;
            font-weight: 700;
            border: 1px solid rgba(255, 255, 255, 0.8);
            box-shadow: 
                0 4px 10px rgba(102, 45, 140, 0.08),
                inset 0 1px 0 rgba(255, 255, 255, 0.9),
                inset 0 -2px 0 rgba(0, 0, 0, 0.03);
            transition: all 0.2s ease;
        }
        div.stDownloadButton > button:hover {
            transform: translateY(-2px);
            background: linear-gradient(145deg, #ffffff 0%, #f0e6ff 100%);
            border-color: #ED1E79;
            box-shadow: 0 8px 15px rgba(102, 45, 140, 0.15);
            color: #ED1E79 !important;
        }
        div[data-baseweb="notification"] {
            background-color: rgba(102, 45, 140, 0.05);
            border-left-color: #662D8C;
            border-radius: 12px;
        }
        </style>
    """, unsafe_allow_html=True)

    st.markdown('<div class="main-title">What can I help with?</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="sub-title">请分别上传您的【周期性复盘报告】、【行业benchmark】数据文件，我将为您生成专业准确的广告优化【数据终表】。</div>', 
        unsafe_allow_html=True
    )

    col1, col_gap, col2 = st.columns([1, 0.1, 1])

    with col1:
        with st.container(border=True):
            st.markdown('<div class="icon-container">📊</div>', unsafe_allow_html=True)
            st.markdown('<div class="card-header">1.上传【周期性复盘报告】</div>', unsafe_allow_html=True)
            raw_files = st.file_uploader(
                "",
                type=["xlsx", "xls"],
                key="raw_uploader",
                label_visibility="collapsed",
                accept_multiple_files=True
            )

    with col2:
        with st.container(border=True):
            st.markdown('<div class="icon-container">🎯</div>', unsafe_allow_html=True)
            st.markdown('<div class="card-header">2.上传【行业 Benchmark]】</div>', unsafe_allow_html=True)
            bench_file = st.file_uploader("", type=["xlsx", "xls"], key="bench_uploader", label_visibility="collapsed")

    st.write("")
    st.write("")

    b_c1, b_c2, b_c3 = st.columns([1, 1, 1])
    with b_c2:
        start_btn = st.button("开始生成数据表 ✦", use_container_width=True)

    if start_btn:
        if not raw_files:
            st.error("⚠️ 请至少上传 [数据报表] 才能继续！")
            return

        processor = AdReportProcessor(raw_files, bench_file)

        try:
            with st.spinner("阶段 1/2: 数据清洗、Top10截断、降维合并..."):
                processor.process_etl()
                st.toast("✅ 阶段 1 完成：Master Tables 已生成", icon="✅")

            with st.expander("📄 点击查看处理后的数据预览 (Master Tables)", expanded=False):
                tabs = st.tabs(list(processor.merged_dfs.keys()))
                for i, (k, v) in enumerate(processor.merged_dfs.items()):
                    with tabs[i]: 
                        st.dataframe(v.head(20), use_container_width=True)

            with st.spinner("阶段 2/2: 生成架构诊断、Word报告 & JSON..."):
                processor.generate_report()
                st.toast("✅ 阶段 2 完成：所有报告已准备就绪", icon="🎉")
            
            st.balloons() 
            
            st.markdown("### 📥 下载结果文件")
            
            with st.container(border=True):
                st.markdown("""
                    <div class="glass-info-box">
                        <span style="font-size: 1.2rem; margin-right: 0.8rem;">💡</span>
                        <span style="
                            font-weight: 600;
                            background: linear-gradient(135deg, #662D8C 0%, #ED1E79 100%);
                            -webkit-background-clip: text;
                            -webkit-text-fill-color: transparent;
                            font-family: -apple-system, BlinkMacSystemFont, sans-serif;
                        ">
                            建议：您可只选择下载 JSON 格式文件用于大模型分析，如有必要再下载其他格式文件。
                        </span>
                    </div>
                """, unsafe_allow_html=True)

                res_c1, res_c2, res_c3 = st.columns(3)

                json_str = json.dumps(processor.final_json, indent=4, ensure_ascii=False)
                res_c1.download_button(
                    "📥 JSON (大模型分析)", 
                    json_str, 
                    "Ad_Report_Data.json", 
                    "application/json",
                    use_container_width=True
                )

                output_xls = io.BytesIO()
                with pd.ExcelWriter(output_xls, engine='xlsxwriter') as writer:
                    for name, df in processor.merged_dfs.items(): 
                        df.to_excel(writer, sheet_name=name, index=False)
                res_c2.download_button(
                    "📥 Excel (数据透视)", 
                    output_xls.getvalue(), 
                    "Merged_Ad_Report_Final.xlsx", 
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )

                output_doc = io.BytesIO()
                processor.doc.save(output_doc)
                res_c3.download_button(
                    "📥 Word (数据审查)", 
                    output_doc.getvalue(), 
                    "Ad_Report_Final_V20_10.docx", 
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

        except Exception as e:
            st.error(f"❌ 处理过程中发生错误: {str(e)}")
            st.exception(e)

if __name__ == "__main__":
    main()
